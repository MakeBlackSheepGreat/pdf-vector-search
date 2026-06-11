#!/usr/bin/env python3
"""
复习资料处理工具 — 从 Word 提取问答，搜索教材页码，生成格式化文档

用法:
    python process_review.py -i 复习资料.docx -o 整理版.docx
    python process_review.py -i 复习资料.docx -o 整理版.docx --title "数据结构"
    python process_review.py -i 复习资料.docx --no-search  # 只整理不搜索
"""

import sys
import os
import re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from config import load_cli_config


# ── 工具函数 ──────────────────────────────────────────────

_QUESTION_RE = re.compile(
    r"^\d+[、，,.．\s]"
    r"|^(试论|试述|简述|论述|怎样|为什么|如何|简答|试分析|"
    r"如何看待|怎样理解|怎样认识|怎样评价|怎样评述|"
    r"试说明|试比较|试解释|请说明|请解释|请分析)"
)
_CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十\d]+[章篇节]")
_HEADING_RE = re.compile(r"^[（(]?\d+[）)]|^[一二三四五六七八九十]+[、，]")


def _is_heading_line(text: str) -> bool:
    """判断文本是否为标题行（数字序号或中文序号）"""
    return bool(_HEADING_RE.match(text)) or text.startswith("第")


def _format_answer_block(doc, lines: list, add_para_fn):
    """将答案文本按段落格式化写入文档"""
    current_block = []
    for line in lines:
        line = line.strip()
        if not line:
            if current_block:
                text = " ".join(current_block)
                add_para_fn(text, bold=_is_heading_line(text))
                current_block = []
        else:
            current_block.append(line)
    if current_block:
        text = " ".join(current_block)
        add_para_fn(text, bold=_is_heading_line(text))


# ── 核心功能 ──────────────────────────────────────────────

def parse_qa_from_docx(docx_path: str) -> list:
    """从 Word 文档中解析问答对

    Returns:
        [{"question": str, "answer": str}, ...]
    """
    from docx import Document

    doc = Document(docx_path)
    raw = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    questions, current_q, current_a = [], None, []

    for text in raw:
        is_q = False
        if _QUESTION_RE.match(text):
            is_q = True
        elif _CHAPTER_RE.match(text):
            is_q = False
        elif "？" in text and len(text) < 80:
            is_q = True

        if is_q:
            if current_q is not None:
                questions.append({"question": current_q, "answer": "\n".join(current_a).strip()})
            current_q, current_a = text, []
        elif current_q is not None:
            current_a.append(text)

    if current_q is not None:
        questions.append({"question": current_q, "answer": "\n".join(current_a).strip()})

    return [q for q in questions if len(q["answer"]) > 20 or "？" in q["question"]]


def search_pages(config, questions: list) -> list:
    """为每道题搜索教材页码（复用 direct_query）"""
    from vector_search import direct_query

    print(f"搜索教材页码（共 {len(questions)} 题）...")
    for i, q in enumerate(questions):
        try:
            results = direct_query(config, q["question"][:50], n_results=2)
            q["pages"] = ", ".join(r["page"] for r in results)
        except Exception:
            q["pages"] = "检索失败"
        print(f"  Q{i+1}: 页码 {q.get('pages', '?')}")

    return questions


def generate_docx(questions: list, output_path: str, title: str = "复习资料",
                  subtitle: str = "整理版", show_pages: bool = True):
    """生成格式化的 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    def _set_run_font(run, font_name, size, bold=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_centered(text, size, bold=False, font_name="黑体", color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p.add_run(text), font_name, size, bold, color)

    def add_para(text, indent=True, bold=False, size=12, font_name="宋体",
                 color=None, first_indent=True):
        p = doc.add_paragraph()
        if indent and first_indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        _set_run_font(p.add_run(text), font_name, size, bold, color)
        return p

    # 标题
    add_centered(f"《{title}》", 24, bold=True)
    add_centered(subtitle, 20, bold=True)
    if show_pages:
        add_centered("（含教材页码标注）", 14, font_name="宋体")
    doc.add_paragraph()
    add_para("本文档基于复习资料整理而成", False, size=11, color=(0x66, 0x66, 0x66))
    doc.add_paragraph()

    # 各题
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        _set_run_font(p.add_run(f"第{i}题"), "黑体", 14, bold=True, color=(0x1A, 0x47, 0x8A))

        add_para(q["question"], indent=False, bold=True, size=13, font_name="黑体")

        if show_pages and q.get("pages"):
            add_para(f"【教材参考页码】第{q['pages']}页",
                     indent=False, size=11, color=(0x88, 0x88, 0x88))

        if not q["answer"]:
            add_para("（待补充）", size=11, color=(0xAA, 0xAA, 0xAA))
        else:
            _format_answer_block(doc, q["answer"].split("\n"), add_para)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p.add_run("—" * 25), "宋体", 10, color=(0xCC, 0xCC, 0xCC))

    # 统计
    doc.add_page_break()
    total_chars = sum(len(q["answer"]) for q in questions)
    add_centered("统计信息", 18, bold=True)
    add_para(f"共 {len(questions)} 道题，约 {total_chars} 字", size=11, color=(0x99, 0x99, 0x99))

    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"\n✅ 文档已生成: {output_path}")
    print(f"   共 {len(questions)} 道题，约 {total_chars} 字")


def main():
    def extra_args(parser):
        parser.add_argument("-i", "--input", required=True, help="输入 .docx 文件")
        parser.add_argument("-o", "--output", default="output/复习资料_整理版.docx", help="输出路径")
        parser.add_argument("--title", default="复习资料", help="文档标题")
        parser.add_argument("--subtitle", default="整理版", help="副标题")
        parser.add_argument("--no-search", action="store_true", help="跳过页码搜索")

    cfg, args = load_cli_config(description="复习资料处理", extra_args_fn=extra_args)

    if not os.path.exists(args.input):
        print(f"❌ 输入文件不存在: {args.input}")
        sys.exit(1)

    print(f"📖 读取: {args.input}")
    questions = parse_qa_from_docx(args.input)
    print(f"✓ 解析出 {len(questions)} 道题")

    if not questions:
        print("⚠  未解析到任何问答")
        sys.exit(1)

    if not args.no_search:
        cfg.validate()
        questions = search_pages(cfg, questions)
    else:
        for q in questions:
            q["pages"] = ""

    generate_docx(questions, args.output, title=args.title,
                  subtitle=args.subtitle, show_pages=not args.no_search)


if __name__ == "__main__":
    main()
